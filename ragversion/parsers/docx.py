"""DOCX document parser."""

import asyncio

try:
    from docx import Document
except ImportError:
    Document = None

from ragversion.exceptions import ParsingError
from ragversion.parsers.base import BaseParser


class DOCXParser(BaseParser):
    """Parser for DOCX documents."""

    def __init__(self):
        if Document is None:
            raise ImportError(
                "python-docx is required for DOCX parsing. Install with: pip install ragversion[parsers]"
            )

    async def parse(self, file_path: str) -> str:
        """
        Parse a DOCX document and extract text content.

        Args:
            file_path: Path to the DOCX file

        Returns:
            Extracted text content

        Raises:
            ParsingError: If parsing fails
        """
        try:
            # Run blocking I/O in executor
            loop = asyncio.get_event_loop()
            text = await loop.run_in_executor(None, self._parse_sync, file_path)
            return text
        except Exception as e:
            raise ParsingError(file_path, e)

    def _parse_sync(self, file_path: str) -> str:
        """Synchronous DOCX parsing."""
        try:
            doc = Document(file_path)
            text_parts = []

            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text:
                    text_parts.append(paragraph.text)

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text:
                            text_parts.append(cell.text)

            return "\n".join(text_parts)
        except Exception as e:
            raise ParsingError(file_path, e)

    def supports(self, file_extension: str) -> bool:
        """Check if this parser supports the given file extension."""
        return file_extension.lower() in [".docx", ".doc"]
